"""
Serviço de processamento de modelos de documentos
Extrai placeholders, campos e metadados de DOCX e PDF
"""
import re
import os
import json
from typing import Dict, List, Any, Optional
from dataclasses import dataclass


@dataclass
class ExtracaoResult:
    sucesso: bool
    campos: List[str]
    placeholders: List[str]
    texto_amostra: str
    tipo_arquivo: str
    erro: Optional[str] = None


def extrair_placeholders_docx(caminho_arquivo: str) -> ExtracaoResult:
    """
    Extrai placeholders no formato {{campo}} de arquivos DOCX
    """
    try:
        from docx import Document
        
        doc = Document(caminho_arquivo)
        texto_completo = ""
        
        # Extrai texto de todos os parágrafos
        for para in doc.paragraphs:
            texto_completo += para.text + "\n"
        
        # Extrai texto de tabelas
        for tabela in doc.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    texto_completo += celula.text + " "
        
        # Encontra placeholders no formato {{campo}}
        padrao = r'\{\{([^}]+)\}\}'
        placeholders = re.findall(padrao, texto_completo)
        
        # Remove duplicatas mantendo ordem
        placeholders_unicos = list(dict.fromkeys(placeholders))
        
        # Limpa e normaliza nomes dos campos
        campos = [p.strip().lower().replace(' ', '_') for p in placeholders_unicos]
        
        return ExtracaoResult(
            sucesso=True,
            campos=campos,
            placeholders=placeholders_unicos,
            texto_amostra=texto_completo[:500] + "..." if len(texto_completo) > 500 else texto_completo,
            tipo_arquivo="docx"
        )
        
    except ImportError:
        return ExtracaoResult(
            sucesso=False,
            campos=[],
            placeholders=[],
            texto_amostra="",
            tipo_arquivo="docx",
            erro="Biblioteca python-docx não instalada"
        )
    except Exception as e:
        return ExtracaoResult(
            sucesso=False,
            campos=[],
            placeholders=[],
            texto_amostra="",
            tipo_arquivo="docx",
            erro=str(e)
        )


def extrair_texto_pdf(caminho_arquivo: str) -> ExtracaoResult:
    """
    Extrai texto de PDF e tenta identificar padrões de campos
    """
    try:
        import PyPDF2
        
        with open(caminho_arquivo, 'rb') as f:
            pdf = PyPDF2.PdfReader(f)
            texto_completo = ""
            
            for pagina in pdf.pages:
                texto_completo += pagina.extract_text() + "\n"
        
        # Tenta identificar padrões comuns de campos em documentos
        padroes = [
            r'\[([^\]]+)\]',  # [campo]
            r'\{\{([^}]+)\}\}',  # {{campo}}
            r'___+',  # Campos sublinhados
            r'\b(NOME|CPF|RG|ENDEREÇO|DATA|ASSINATURA)\b[:\s]*[\n\r]',  # Labels comuns
        ]
        
        campos_encontrados = []
        for padrao in padroes:
            matches = re.findall(padrao, texto_completo, re.IGNORECASE)
            campos_encontrados.extend(matches)
        
        # Remove duplicatas
        campos_unicos = list(dict.fromkeys([c.strip() for c in campos_encontrados if len(c.strip()) > 2]))
        
        return ExtracaoResult(
            sucesso=True,
            campos=campos_unicos,
            placeholders=campos_unicos,
            texto_amostra=texto_completo[:500] + "..." if len(texto_completo) > 500 else texto_completo,
            tipo_arquivo="pdf"
        )
        
    except ImportError:
        return ExtracaoResult(
            sucesso=False,
            campos=[],
            placeholders=[],
            texto_amostra="",
            tipo_arquivo="pdf",
            erro="Biblioteca PyPDF2 não instalada"
        )
    except Exception as e:
        return ExtracaoResult(
            sucesso=False,
            campos=[],
            placeholders=[],
            texto_amostra="",
            tipo_arquivo="pdf",
            erro=str(e)
        )


def processar_modelo(caminho_arquivo: str) -> Dict[str, Any]:
    """
    Processa arquivo de modelo e retorna metadados extraídos
    """
    extensao = os.path.splitext(caminho_arquivo)[1].lower()
    
    if extensao == '.docx':
        resultado = extrair_placeholders_docx(caminho_arquivo)
    elif extensao == '.pdf':
        resultado = extrair_texto_pdf(caminho_arquivo)
    else:
        return {
            "sucesso": False,
            "erro": f"Tipo de arquivo não suportado: {extensao}",
            "campos": [],
            "placeholders": [],
            "tipo_arquivo": extensao
        }
    
    return {
        "sucesso": resultado.sucesso,
        "erro": resultado.erro,
        "campos": resultado.campos,
        "placeholders": resultado.placeholders,
        "texto_amostra": resultado.texto_amostra,
        "tipo_arquivo": resultado.tipo_arquivo,
        "total_campos": len(resultado.campos)
    }


def gerar_preview_preenchido(caminho_modelo: str, dados: Dict[str, str]) -> str:
    """
    Gera preview do documento com dados preenchidos (substitui placeholders)
    """
    try:
        extensao = os.path.splitext(caminho_modelo)[1].lower()
        
        if extensao == '.docx':
            from docx import Document
            
            doc = Document(caminho_modelo)
            
            # Substitui em parágrafos
            for para in doc.paragraphs:
                for campo, valor in dados.items():
                    placeholder = f"{{{{{campo}}}}}"
                    if placeholder in para.text:
                        para.text = para.text.replace(placeholder, str(valor))
            
            # Substitui em tabelas
            for tabela in doc.tables:
                for linha in tabela.rows:
                    for celula in linha.cells:
                        for campo, valor in dados.items():
                            placeholder = f"{{{{{campo}}}}}"
                            if placeholder in celula.text:
                                celula.text = celula.text.replace(placeholder, str(valor))
            
            # Salva preview temporário
            preview_path = caminho_modelo.replace('.docx', '_preview.docx')
            doc.save(preview_path)
            
            return preview_path
            
        else:
            return "Preview disponível apenas para DOCX"
            
    except Exception as e:
        return f"Erro ao gerar preview: {str(e)}"


def validar_campos_obrigatorios(caminho_modelo: str, dados_fornecidos: Dict[str, str]) -> Dict[str, Any]:
    """
    Valida se todos os campos obrigatórios do modelo foram preenchidos
    """
    metadados = processar_modelo(caminho_modelo)
    
    if not metadados["sucesso"]:
        return {
            "valido": False,
            "erro": metadados.get("erro", "Erro ao processar modelo"),
            "campos_faltando": []
        }
    
    campos_modelo = set(metadados["campos"])
    campos_fornecidos = set(dados_fornecidos.keys())
    
    campos_faltando = list(campos_modelo - campos_fornecidos)
    
    return {
        "valido": len(campos_faltando) == 0,
        "campos_faltando": campos_faltando,
        "campos_preenchidos": list(campos_fornecidos),
        "total_campos": len(campos_modelo)
    }
